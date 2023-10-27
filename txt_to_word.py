from docx import Document
import os

path = "./text/"

file_list = os.listdir(path)

if not os.path.exists("./output"):
    os.makedirs("./output")

for file in file_list:

    filename = os.path.splitext(file)[0]
    document = Document()
    document.add_heading(filename, 0)

    # Get file contents
    # If on Windows, specify encoding="utf-8" for open function
    try:
        file_contents = open(path + file).read()

    # Break loop if file is binary (non-text)
    except UnicodeDecodeError:
        break

    # Remove unwanted words
    final_text = file_contents.replace(" آآ ", " ").replace(
        " اا ", " ").replace(" اه ", " ")

    # Add final text to the word document
    p = document.add_paragraph(final_text)

    document.save('./output/' + filename + '.docx')
